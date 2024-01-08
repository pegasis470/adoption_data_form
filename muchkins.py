from csv import *
from tkinter import *
from datetime import datetime as dt
import datetime
import random
import os
from tkinter import messagebox
import cv2
from PIL import Image, ImageTk
import docx
from docx.shared import Inches,Cm,RGBColor,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import font
#This is the Backend of the application
## global varibale contain csv file name and path 
date=datetime.date.today()
filename=str(date)+'.csv'
data_path=os.path.join(os.getcwd(),'data',filename)

## This is the camara class with all its oprations DO NOT EDIT!!
class Camara:
    def __init__(self,imageof):
        self.imageof=imageof
        self.photo=0
        self.app=0
        self.label_widget=0
    def show_photo(self):
        self.photo=Toplevel()
        label=Label(self.photo)
        opencv_image=cv2.imread(os.path.join('images','Temp_name.png'))
        opencv_image = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2RGBA)
        captured_image = Image.fromarray(opencv_image)
        photo_image = ImageTk.PhotoImage(image=captured_image)
        label.photo_image = photo_image
        label.configure(image=photo_image)
        label.pack()
        Button(self.photo,text='SAVE',command=self.Save_image).pack()
    def Take_photo(self):
        _, frame = self.vid.read()
        opencv_image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB) 
        cv2.imwrite(os.path.join('images','Temp_name.png'), cv2.cvtColor(opencv_image, cv2.COLOR_RGB2BGR))
        self.show_photo()
    def open_camera(self): 
    
        # Capture the self.video frame by frame 
        _, frame = self.vid.read() 
      
        # Convert image from one color space to other 
        opencv_image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA) 
      
        # Capture the latest frame and transform to image 
        captured_image = Image.fromarray(opencv_image) 
      
        # Convert captured image to photoimage 
        photo_image = ImageTk.PhotoImage(image=captured_image) 
      
        # Displaying photoimage in the label 
        self.label_widget.photo_image = photo_image 
      
        # Configure image in the label 
        self.label_widget.configure(image=photo_image) 
        # Repeat the same process after every 10 seconds 
        self.label_widget.after(10, self.open_camera) 
      
      
    def New_window(self):
        self.app = Toplevel()
        self.label_widget= Label(self.app)
        self.label_widget.pack()
        button2= Button(self.app,text="Take image",command=self.Take_photo)
        button2.pack()
        ## prequisites for the camera 
        self.vid = cv2.VideoCapture(0)
        width, height = 800, 600
        self.vid.set(cv2.CAP_PROP_FRAME_WIDTH, width)
        self.vid.set(cv2.CAP_PROP_FRAME_HEIGHT, height)
        self.open_camera()
    def end(self):
        self.app.destroy()
        self.photo.destroy()
    def Save_image(self):
        print(self.imageof)
        if self.imageof==0:
            os.rename(os.path.join('images','Temp_name.png'),os.path.join('images',f'{Tag.get()}.png'))
        elif self.imageof==1:
            os.rename(os.path.join('images','Temp_name.png'),os.path.join('images',f'{Caretaker.get()}.png'))
        elif self.imageof==2:
            os.rename(os.path.join('images','Temp_name.png'),os.path.join('images',f'{Adopter.get()}.png'))
        messagebox.showinfo("Sucess","Image captured sucessfully")
        self.vid.release()
        self.end()
## declear a global var to access this class
## class to test and create csv file
class BackEnd:
    def __init__(self):
        date=datetime.date.today()
        filename=str(date)+'.csv'
        self.data_path=os.path.join(os.getcwd(),'data',filename)
        self.file_path=os.path.join(os.getcwd(),'Forms',f'{Adopter.get()}.docx')
        self.Test_file()                
    def Test_file(self):
        try:
            file=open(data_path)
            file.close()
        except FileNotFoundError :
            try:
                open_file=open(data_path,'w')
                Writer=writer(open_file)
                Writer.writerow(['counselor name','Date','Tag number','Sex','Age','Color','Breed','Physical Fitness','Vaccination','Sterlisation','Caretaker/Foster name','caretaker Contact no','caretaker whatsapp','Email','Local Residence','Permanent residence','Adopter\'s name','adopter Contact no','adopter whatsapp','Email','Government ID','ID Number','Social ID','Local Residence','Permanent Residence','What are your plans for your pet if you shift to another town/place?','What are your plans for your pet if you shift to another town/place?','Amount of time your pet may have to be alone in a day?','Who will take care of your pet if you go out of town temporarily?'])
                open_file.close()
            except FileNotFoundError:
                os.mkdir(os.path.join(os.getcwd(),'data'))
                Test_file()
    def Check_duplicate(self):
        path=os.path.join(os.getcwd(),'Forms')
        file_name =f"{Adopter.get()}.docx"
        try:
            files=os.listdir(path)
        except FileNotFoundError:
            os.mkdir(path)
            files=os.listdir(path)
        if file_name in files:
            messagebox.showerror("ERROR","Two files with same name found")
            return False
        else:
            return True


    def Write_form(self):
        image_animal=os.path.join('images',f'{Tag.get()}.png')
        caretaker_image=os.path.join('images',f'{Caretaker.get()}.png')
        adopter_image=os.path.join('images',f'{Adopter.get()}.png')
        doc = docx.Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
        doc.add_paragraph("Serial NO - __________________")
        doc.add_picture('logo.jpg', width=Inches(5), height=Inches(1.5))
        img=doc.paragraphs[-1]
        img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1=doc.add_heading("ADOPTION AND CONSENT FORM",0)
        h1.style.font.color.rgb = RGBColor(0,0,0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Councler Name: {Councler.get()}					Date: {datetime.date.today()} ")
        dic=doc.add_heading("DESCRIPTION OF ANIMAL",1)
        dic.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dic.style.font.color.rgb = RGBColor(0,0,0)
        animal=doc.add_table(rows=1,cols=2) 
        animal.cell(0,0).add_paragraph(f"Colour: {Color.get()}        Tag:{Tag.get()}\nGender: {Gender.get()}    Age: {Age.get()}\nBreed: {Breed.get()}\nPhysically Fitness Status: {Health.get()}\nVaccination Status: {Vaccination.get()}\nSterilisation Status: {Sterilisation.get()}")
        animal.cell(0,1).add_paragraph().add_run().add_picture(image_animal,width=Inches(2),height=Inches(1.75))
        details=doc.add_heading("DETAILS OF CARETAKER                   DETAILS OF ADOPTER",1)
        details.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details.style.font.color.rgb = RGBColor(0,0,0)
        caretaker=doc.add_table(rows=2,cols=2)
        if SameVar.get():
            formated_text=f"Name of Caretaker: {Caretaker.get()}\nContact No.: {CaretakerContact.get()}\nWhatsapp No.: {CaretakerWhatsapp.get()}\nEmail I'd: {Email.get()}\nLocal Residence: {LocalAdd.get()}\nPermanent Residence: {LocalAdd.get()}\nInstagram/Facebook ID: {ID.get()}\nSigneture:_______________"
        else:
            formated_text=f"Name of Caretaker: {Caretaker.get()}\nContact No.: {CaretakerContact.get()}\nWhatsapp No.: {CaretakerWhatsapp.get()}\nEmail I'd: {Email.get()}\nLocal Residence: {LocalAdd.get()}\nPermanent Residence: {PermanentAdd.get()}\nInstagram/Facebook ID: {ID.get()}\nSigneture:_______________"
        caretaker.cell(0,0).add_paragraph(formated_text)
        if SameVar1.get():
            formated_text=f"Name of Adopter: {Adopter.get()}\nContact No.: {AdopterContact.get()}\nWhatsapp No.: {AdopterWhatsapp.get()}\nEmail Id: {Email1.get()}      Government ID:{gov_ID.get()}\nLocal Residence: {AdopterLocalAdd.get()}          ID No.:{ID_num.get()}\nPermanent Residence: {AdopterLocalAdd.get()}\nInstagram/Facebook ID: {Adoptor_ID.get()}\nSigneture:_______________" 
        else:
           formated_text=f"Name of Adopter: {Adopter.get()}\nContact No.: {AdopterContact.get()}\nWhatsapp No.: {AdopterWhatsapp.get()}\nEmail Id: {Email1.get()}\nLocal Residence: {AdopterLocalAdd.get()}\nPermanent Residence: {AdopterPermanentAdd.get()}\nInstagram/Facebook ID: {Adoptor_ID.get()}\nSigneture:_______________" 
        caretaker.cell(0,1).add_paragraph(formated_text)
        caretaker.cell(1,0).add_paragraph().add_run().add_picture(caretaker_image,width=Inches(2),height=Inches(2))
        img=caretaker.cell(1,1).add_paragraph()
        img.add_run().add_picture(adopter_image,width=Inches(2),height=Inches(2))
        img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading("General information from Adopter",1)
        para=doc.add_paragraph(f"What are your plans for your pet if you shift to another town/place?\n {Plans.get('1.0','end-1c')}")
        para.add_run(f"\nHave you had a pet before or have one right now?\n{Owned.get()}")
        para.add_run(f"\nAmount of time your pet may have to be alone in a day?\n{Alone.get()}")
        para.add_run(f"\nWho will take care of your pet if you go out of town temporarily?\n{NotHome.get()}")
        con=doc.add_heading("Consent By Adopter",1)
        con.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        con.style.font.color.rgb = RGBColor(0,0,0)
        doc.add_paragraph('With the unanimous consent of all my family members, I am willingly adopting the pet, assuming full responsibility and acknowledging that I will always regard my pet as an integral part of our family, ensuring that it is never treated merely as an object.',style='List Bullet')
        doc.add_paragraph("I commit to maintaining a clean and well-ventilated environment for the adopted pet, providing proper nourishment, and ensuring regular exercise. I will refrain from keeping the animal tethered or chained with an unreasonably short or heavy restraint for an extended duration.",style='List Bullet')
        doc.add_paragraph("I acknowledge that certain diseases may not be detectable during the initial checkup of the adopted pet, and the Animals With Humanity Team will not be held responsible for such conditions. Therefore, it is recommended to conduct a comprehensive post-adoption health examination.",style='List Bullet')
        doc.add_paragraph("If the pet I adopt becomes unwell, I will promptly seek advice from a veterinarian and notify the Animals With Humanity Team. Additionally, I will take responsibility for adhering to the deworming, vaccination, and sterilization schedule for the well-being of the pet.",style='List Bullet')
        doc.add_paragraph("Should the responsibility of pet parenting need to be transferred, I will provide advance notice to both the caretaker and Team Animals With Humanity. I understand that the adoption process will need to be repeated for the new caretakers, and a fine of Rs. 5,000/- will be duly acknowledged and paid.",style='List Bullet')
        doc.add_paragraph("I acknowledge that the Animals With Humanity Team possesses the authority to conduct unannounced inspections of the conditions in which I'm taking care for the adopted pet. In the event of any violations, the Animals With Humanity Team is empowered to take legal action as per applicable law.",style='List Bullet')
        doc.add_paragraph("I ______________ acknowledge that abandoning or subjecting my pet to mistreatment may lead to legal consequences under the Prevention of Cruelty to Animals Act of 1960. In case my pet is found in any such situation, Team Animals With Humanity will take a fine upto Rs. 10,000/- depending on the situation and proceed with legal action.",style='List Bullet')
        doc.add_paragraph("I enter into this contract of my own free will and understand that this is a binding contract enforceable by civil law.",style='List Bullet')
        doc.save(self.file_path)
        doc.save(os.path.join(os.getcwd(),'Print.docx'))
    def Append_data(self):
        data=[Councler.get(),dt.now().strftime("%d.%m.%y %I:%m %p"),Tag.get(),Gender.get(),Age.get(),Color.get(),Breed.get(),Health.get(),Vaccination.get(),Sterilisation.get(),Caretaker.get(),CaretakerContact.get(),CaretakerWhatsapp.get(),Email.get(),LocalAdd.get(),PermanentAdd.get(),Adopter.get(),AdopterContact.get(),AdopterWhatsapp.get(),Email1.get(),gov_ID.get(),ID_num,Adoptor_ID.get(),AdopterLocalAdd.get(),PermanentAdd.get(),Plans.get("1.0","end-1c"),Owned.get(),Alone.get(),NotHome.get()]
        try:
                file=open(self.data_path,'a') 
        except PermissionError:
                messagebox.showerror('ERROR',"The data base file is locked as it may be in use by another application please close that instance and try agian ")
                return None
        Writer=writer(file)
        Writer.writerow(data)
        file.close()
        messagebox.showinfo("success","data saved") 
    def clear(self):
        Councler.delete(0,END)
        Gender.delete(0,END)
        Age.delete(0,END)
        Color.delete(0,END)
        Breed.delete(0,END)
        Tag.delete(0,END)
        Health.delete(0,END)
        Vaccination.delete(0,END)
        Sterilisation.delete(0,END)
        Caretaker.delete(0,END)
        CaretakerContact.delete(0,END)
        CaretakerWhatsapp.delete(0,END)
        Email.delete(0,END)
        ID.delete(0,END)
        LocalAdd.delete(0,END)
        PermanentAdd.delete(0,END)
        Adopter.delete(0,END)
        AdopterContact.delete(0,END)
        AdopterWhatsapp.delete(0,END)
        Adoptor_ID.delete(0,END)
        Email1.delete(0,END)
        gov_ID.delete(0,END)
        ID_num.delete(0,END)
        AdopterLocalAdd.delete(0,END)
        PermanentAdd.delete(0,END)
        Plans.delete("1.0","end")
        Owned.delete(0,END)
        Alone.delete(0,END)
        NotHome.delete(0,END)
def call_backend():
    back=BackEnd()
    back.Write_form()
    back.Append_data()
    back.clear()
    os.startfile("Print.docx", "print")
# This is the front end of the application
##function to show the hindi version of T&C
def translate():
    new=Toplevel()
    new.iconphoto(False,icon)
    new.title("नियम शर्तें और सहमति प्रपत्र")
    Label(new,text="एडॉप्टर द्वारा सहमति",font=("Arial",20,"bold")).pack()
    Label(new,text="1)  मेरे सभी परिवार के सदस्यों की सर्वसम्मति से, मैं स्वेच्छा से पालतू जानवर को गोद ले रहा हूं, पूरी जिम्मेदारी संभालने और यह स्वीकार करते हुए कि मैं हमेशा अपने पालतू को अपने परिवार का एक अभिन्न",font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text="    अंग मानूंगा, यह सुनिश्चित करते हुए कि इसे कभी भी केवल एक वस्तु के रूप में नहीं माना जाता है.",font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text= "2) मैं गोद लिए गए पालतू जानवरों के लिए एक स्वच्छ और अच्छी तरह हवादार वातावरण बनाए रखने, उचित पोषण प्रदान करने और नियमित व्यायाम सुनिश्चित करने के लिए प्रतिबद्ध हूं.",font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text= "   मैं जानवरों को एक विस्तारित अवधि के लिए अनुचित रूप से कम या भारी संयम के साथ रखने या जंजीर रखने से बचना चाहूंगा.",font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='3)  मैं स्वीकार करता हूं कि गोद लिए गए पालतू जानवरों की प्रारंभिक जांच के दौरान कुछ बीमारियों का पता नहीं लगाया जा सकता है, और पशु के साथ मानवता टीम को',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    ऐसी स्थितियों के लिए जिम्मेदार नहीं ठहराया जाएगा. इसलिए, एक व्यापक पोस्ट-गोद लेने वाली स्वास्थ्य परीक्षा आयोजित करने की सिफारिश की जाती है.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='4)  अगर मैं जिस पालतू जानवर को गोद लेता हूं वह अस्वस्थ हो जाता है, तो मैं तुरंत एक पशुचिकित्सा से सलाह लूंगा और जानवरों को मानवता टीम के साथ सूचित करूंगा.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    इसके अतिरिक्त, मैं पालतू जानवरों की भलाई के लिए ओसिंग, टीकाकरण और नसबंदी अनुसूची का पालन करने की जिम्मेदारी लूंगा.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='5)  क्या पालतू पालन-पोषण की जिम्मेदारी को हस्तांतरित करने की आवश्यकता है, मैं कार्यवाहक और टीम एनिमल्स दोनों को मानवता के लिए अग्रिम सूचना प्रदान करूंगा.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    मैं समझता हूं कि नए कार्यवाहकों के लिए गोद लेने की प्रक्रिया को दोहराया जाना चाहिए, और रुपये का जुर्माना। 5,000 / - विधिवत स्वीकार और भुगतान किया जाएगा.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='6)  मैं स्वीकार करता हूं कि एनिमल्स विद ह्यूमैनिटी टीम के पास उन परिस्थितियों का अघोषित निरीक्षण करने का अधिकार है, जिनमें मैं ',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    गोद लिए गए पालतू जानवरों की देखभाल कर रहा हूं. किसी भी उल्लंघन की स्थिति में, पशु के साथ मानवता टीम को लागू कानून के अनुसार कानूनी कार्रवाई करने का अधिकार है.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='7)  मैं (दत्तक) \ स्वीकार करता हूं कि मेरे पालतू जानवर को दुर्व्यवहार के लिए छोड़ने या अधीन करने से 1960 के पशु अधिनियम के क्रूरता निवारण',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    के तहत कानूनी परिणाम हो सकते हैं. यदि मेरा पालतू ऐसी किसी भी स्थिति में पाया जाता है, तो टीम एनिमल्स विद ह्यूमैनिटी रुपये तक का जुर्माना लेगी।',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    10,000 / - स्थिति के आधार पर और कानूनी कार्रवाई के साथ आगे बढ़ें.',font=T_C_font).pack(pady=2,anchor="w")
    Label(new,text='    मैं अपनी मर्जी के इस अनुबंध में प्रवेश करता हूं और समझता हूं कि यह नागरिक कानून द्वारा लागू एक बाध्यकारी अनुबंध है.',font=T_C_font).pack()
    Button(new,text="ठीक",command=new.destroy).pack()


## primary
window=Tk()
icon=PhotoImage(file=os.path.join(os.getcwd(),'icon.png'))
window.iconphoto(False,icon)
window.title('ADOPTION FORM')
#window.geometry('2')
window.protocol("WM_DELETE_WINDOW", window.destroy)
Label(window,text="Counselor name").grid(sticky="E",row=0,column=0)
Councler=Entry(window)
Councler.grid(sticky="W",row=0,column=1)
Label(window,text="Description of the Animal:",font=("Arial",15,"bold","underline")).grid(sticky="W",row=1,column=1)
#sex.grid(sticky="W",row=2,column=1)
Label(window,text="Tag Number").grid(sticky="E",row=2,column=0)
Tag=Entry(window)
Tag.grid(sticky="W",row=2,column=1)
Label(window,text='Gender').grid(sticky="E",row=3,column=0)
Gender=Entry(window)
Gender.grid(sticky="W",row=3,column=1)
Label(window,text='Age').grid(sticky="E",row=3,column=2)
Age=Entry(window)
Age.grid(sticky="W",row=3,column=3)
Label(window,text='Colour').grid(sticky="E",row=4,column=0)
Color=Entry(window)
Color.grid(sticky="W",row=4,column=1)
Label(window,text='Breed').grid(sticky="E",row=4,column=2)
Breed=Entry(window)
#Lable(window,text='Tag Number').grid(sticky="E",row=3,column=2)
Breed.grid(sticky="W",row=4,column=3)
Label(window,text='Physically Fitness Status').grid(sticky="E",row=5,column=0)
Health=Entry(window)
Health.grid(sticky="W",row=5,column=1)
Label(window,text="Vaccination status").grid(sticky="E",row=5,column=2)
Vaccination=Entry(window)
Vaccination.grid(sticky="E",row=5,column=3)
Label(window,text="Sterilisation Status").grid(sticky="E",row=6,column=0)
Sterilisation=Entry(window)
Sterilisation.grid(sticky="W",row=6,column=1)
Button(window,text="Photo Of Animal",command=Camara(0).New_window).grid(row=6,column=2)
Label(window,text="Specifics of the Caretaker/Foster",font=("Arial",15,"bold","underline")).grid(sticky="W",row=7,column=1)
Label(window,text='Caretaker/Foster’s name').grid(sticky="E",row=8,column=0)
Caretaker=Entry(window)
Caretaker.grid(sticky="W",row=8,column=1)
Label(window,text='Contact no').grid(sticky="E",row=9,column=0)
CaretakerContact=Entry(window)
CaretakerContact.grid(sticky="W",row=9,column=1)
Label(window,text='Contact no/Whatsapp').grid(sticky="E",row=9,column=2)
CaretakerWhatsapp=Entry(window)
CaretakerWhatsapp.grid(sticky="E",row=9,column=3)
Label(window,text='Email').grid(sticky="E",row=10,column=0)
Email=Entry(window)
Email.grid(sticky="W",row=10,column=1)
Label(window,text='Instagram/Facebook ID').grid(sticky="E",row=10,column=2)
ID=Entry(window)
ID.grid(sticky="W",row=10,column=3)
Label(window,text='Local Residence').grid(sticky="E",row=11,column=0)
LocalAdd=Entry(window)
LocalAdd.grid(sticky="W",row=11,column=1)
Button(window,text="Photo of Caretaker",command=Camara(1).New_window).grid(row=11,column=2)
SameVar = IntVar()
Same1=Checkbutton(window,text="same as local",variable=SameVar,onvalue=True,offvalue=False)
### handle this in backend call SameVar.get() at time of submitting ##
Same1.grid(sticky="W",row=12,column=2)
Label(window,text='Permanent Residence').grid(sticky="E",row=12,column=0)
PermanentAdd=Entry(window)
PermanentAdd.grid(sticky="W",row=12,column=1)
Label(window,text="Specifics of the Adopter",font=("Arial",15,"bold","underline")).grid(sticky="W",row=13,column=1)
Label(window,text="Adopter's Name").grid(sticky="E",row=14,column=0)
Adopter=Entry(window)
Adopter.grid(sticky="W",row=14,column=1)
Label(window,text='Contact No').grid(sticky="E",row=15,column=0)
AdopterContact=Entry(window)
AdopterContact.grid(sticky="W",row=15,column=1)
Label(window,text='Contact No/Whatsapp').grid(sticky="E",row=15,column=2)
AdopterWhatsapp=Entry(window)
AdopterWhatsapp.grid(sticky="W",row=15,column=3)
Label(window,text='Instagram/Facebook ID').grid(sticky="E",row=16,column=0)
Adoptor_ID=Entry(window)
Adoptor_ID.grid(sticky="W",row=16,column=1)
Button(window,text="Photo of Adopter",command=Camara(2).New_window).grid(row=16,column=2)
Label(window,text='Email').grid(sticky="E",row=17,column=0)
Email1=Entry(window)
Email1.grid(sticky="W",row=17,column=1)
Label(window,text="Government ID").grid(sticky="E",row=17,column=2)
gov_ID=Entry(window)
gov_ID.grid(sticky="W",row=17,column=3)
Label(window,text="ID Number").grid(sticky="E",row=18,column=2)
ID_num=Entry(window)
ID_num.grid(sticky="W",row=18,column=3)
Label(window,text='Local Address').grid(sticky="E",row=18,column=0)
AdopterLocalAdd=Entry(window)
AdopterLocalAdd.grid(sticky="W",row=18,column=1)
SameVar1 = IntVar()
Same2=Checkbutton(window,text="same as local",variable=SameVar1,onvalue=True,offvalue=False)
### handle this in backend call SameVar.get() at time of submitting ##
Same2.grid(sticky="W",row=19,column=2)
Label(window,text='Permanent Residence').grid(sticky="E",row=19,column=0)
AdopterPermanentAdd=Entry(window)
AdopterPermanentAdd.grid(sticky="W",row=19,column=1)
Label(window,text='What are your plans for your pet if you shift to another town/place?').grid(sticky="e",row=20,column=0,columnspan=2)
Plans=Text(window,height=3,width=45)
Plans.grid(sticky="W",row=20,column=2,columnspan=2)
Label(window,text='Have you had a pet before or have one right now?').grid(sticky="E",row=21,column=0)
Owned=Entry(window)
Owned.grid(sticky="W",row=21,column=1)
Label(window,text='Amount of time your pet may have to be alone in a day?').grid(sticky="E",row=22,column=0)
Alone=Entry(window)
Alone.grid(sticky="W",row=22,column=1)
Label(window,text='Who will take care of your pet if you go out of town temporarily?').grid(sticky="W",row=23,column=0)
NotHome=Entry(window)
NotHome.grid(sticky="W",row=23,column=1)
button=Button(window,text="save,print and clear",command=call_backend)
button.grid(row=24,column=1)
backend=BackEnd()
backend.Test_file()

##T&C window
Terms_window=Toplevel()
T_C_font=font.Font(weight="bold")
Terms_window.iconphoto(False,icon)
Terms_window.title('Terms and Conditions and Consent form')
Label(Terms_window,text='CONSENT BY ADOPTER',font=("Arial", 25,"bold")).pack()
Label(Terms_window,text='1) With the unanimous consent of all my family members, I am willingly adopting the pet, assuming full responsibility and acknowledging that',font=T_C_font).pack(pady=2,anchor="w") 
Label(Terms_window,text='    I will always regard my pet as an integral part of our family, ensuring that it is never treated merely as an object.',font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="2) I commit to maintaining a clean and well-ventilated environment for the adopted pet, providing proper nourishment, and ensuring regular",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    exercise. I will refrain from keeping the animal tethered or chained with an unreasonably short or heavy restraint for an extended duration.",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="3) I acknowledge that certain diseases may not be detectable during the initial checkup of the adopted pet, and the Animals With Humanity Team",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    will not be held responsible for such conditions. Therefore, it is recommended to conduct a comprehensive post-adoption health examination.",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="4) If the pet I adopt becomes unwell, I will promptly seek advice from a veterinarian and notify the Animals With Humanity Team. Additionally,",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    I will take responsibility for adhering to the deworming, vaccination, and sterilization schedule for the well-being of the pet.",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="5) Should the responsibility of pet parenting need to be transferred, I will provide advance notice to both the caretaker and",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    Team Animals With Humanity. I understand that the adoption process will need to be repeated for the new caretakers,",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    and a fine of Rs. 5,000/- will be duly acknowledged and paid.",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="6) I acknowledge that the Animals With Humanity Team possesses the authority to conduct unannounced inspections of the",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    conditions in which I'm taking care for the adopted pet. In the event of any violations, the Animals With Humanity Team is",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    empowered to take legal action as per applicable law.",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="7) I (Adopter) acknowledge that abandoning or subjecting my pet to mistreatment may lead to legal",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    consequences under the Prevention of Cruelty to Animals Act of 1960. In case my pet is found in any such situation,",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="    Team Animals With Humanity will take a fine upto Rs. 10,000/- depending on the situation and proceed with legal action.",font=T_C_font).pack(pady=2,anchor="w")
Label(Terms_window,text="I enter into this contract of my own free will and understand that this is a binding contract enforceable by civil law.",font=T_C_font).pack()
check=Checkbutton(Terms_window,text='I consent to the above terms',font=T_C_font)
check.pack()
Button(Terms_window,text='translate to hindi',command=translate).pack()
Terms_window.mainloop()

if __name__=='__main__':
    ##preform backend prequisites here## 
    window.mainloop()

