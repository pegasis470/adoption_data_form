from csv import *
from tkinter import *
import tkinter as Top
from tkinter import messagebox
from datetime import datetime as dt
import datetime
import random
import os
import docx 




date=datetime.date.today()
filename=str(date)+'.csv'
path=os.path.join(os.getcwd(),'data',filename)
columns=['counselor name','Date','Time','Kind of Animal','Tag number','Sex','Age','Fur Color','Breed','Health Acknowledgement','Caretaker/Foster name','Caretaker Contact no','Caretaker whatsapp','Email','Local Residence','Permanent residence','Adopter\'s name','Adoptor Contact no','Adoptor whastapp','Line of work','Email','Local Residence','Permanent Residence','What are your plans for taking care of your','Have you had a pet before','Your pet will primarly be an','My dog/cat needs to be able to be alone (per day)','When not home']
def Test_file():
    try:
        file=open(path)
        file.close()
    except FileNotFoundError :
        try:
            open_file=open(path,'w')
            Writer=writer(open_file)
            Writer.writerow(['counselor name','Date','Time','Kind of Animal','Tag number','Sex','Age','Color','Breed','Health Acknowledgement','Caretaker/Foster name','caretaker Contact no','caretaker whatsapp','Email','Local Residence','Permanent residence','Adopter\'s name','adoptor Contact no','adoptor whastapp','Line of work','Email','Local Residence','Permanent Residence','incase of moving','pet before or currently','Your pet will primarly be an','to be alone (per day)','when not home'])
            open_file.close()
        except FileNotFoundError:
            os.mkdir(os.path.join(os.getcwd(),'data'))
            Test_file()

def Check_duplicate(path,file_name):
    try:
        files=os.listdir(path)
    except FileNotFoundError:
        return True
    if file_name in files:
        messagebox.showerror("ERROR","two files with same name found a random intiger will be added and the end of this file name")
        return False
    else:
        return True		


def write_text(document,text,head='t'):
    para=document.add_paragraph().add_run()
    if head=='t':
        para.add_text(text)
    else:
        document.add_heading(text)

      
def Save():
    data=[entry7.get(),date,dt.now().strftime('%I:%M %p'),entry1.get(),entry2.get(),entry22.get(),entry3.get(),entry4.get(),entry5.get(),entry6.get(),entry8.get(),entry9.get(),entry10.get(),entry11.get(),entry12.get(),entry13.get(),entry14.get(),entry15.get(),entry16.get(),entry17.get(),entry18.get(),entry19.get(),entry20.get(),entry21.get(),menu2.get(),menu3.get(),menu4.get(),menu5.get()]
    
    try:
            file=open(path,'a')
    except PermissionError:
            messagebox.showerror('ERROR',"The data base file is locked as it may be in use by another application please close that instance and try agian ")
            return None
    Writer=writer(file)
    Writer.writerow(data)
    file.close()
    #print_path=os.path.join('~',media,<device name>,'Prints',f'{entry14.get()}.docx')
    print_path=os.path.join(os.getcwd(),'Prints')
    file_name=f"{entry14.get()}.docx"
    while True:
        if Check_duplicate(print_path,file_name):
            try:
                final_path=os.path.join(print_path,file_name)
                print_file=open(final_path,'w')
            except FileNotFoundError:
                os.mkdir(os.path.join(os.getcwd(),"Prints"))
            document=docx.Document()
            document.add_picture("logo.png")
            for i in range(len(data)):
                if i == 2:
                    write_text(document,"\n Description of the Animal \n",head='h')
                elif i == 9:
                    write_text(document,"\n Specifics of the Caretaker/Foster \n",head='h')
                elif i== 16:
                    write_text(document,"\n Specifics of the Adopter \n",head='h')
                write_text(document,f"{columns[i]}:    {data[i]}\n")
            write_text(document,f"\n\n\nI ({entry14.get()}) has understood and agreee to the tearms and conditions displayed to me on the screen and here by provide my concent by signing below\n\n")
            write_text(document,"adoptior's signeture:________________________________________________________________\n\n")
            write_text(document,"caretaker's signeture:_______________________________________________________________\n\n")
            
            document.save(final_path)
            entry1.delete(0,END),entry2.delete(0,END),entry22.delete(0,END),entry3.delete(0,END),entry4.delete(0,END),entry5.delete(0,END),entry6.delete(0,END),entry7.delete(0,END),entry8.delete(0,END),entry9.delete(0,END),entry10.delete(0,END),entry11.delete(0,END),entry12.delete(0,END),entry13.delete(0,END),entry14.delete(0,END),entry15.delete(0,END),entry16.delete(0,END),entry17.delete(0,END),entry18.delete(0,END),entry19.delete(0,END),entry20.delete(0,END),entry21.delete(0,END),menu2.set('please choose one '),menu3.set('please choose one '),menu4.set('please choose one '),menu5.set('please choose one ')
            check.deselect()
            messagebox.showinfo("sucess",'data saved sucessfully')
            break
        else :
            file_name=f"{entry14.get()}({random.randint(0,100)}).docx"
            continue
def translate():
    new=Toplevel()
    new.iconphoto(False,icon)
    new.title("नियम शर्तें और सहमति प्रपत्र")
    Label(new,text="एडॉप्टर द्वारा सहमति",font=("Arial",20,"bold")).pack()
    Label(new,text="1) मैं पिल्ला/बिल्ली के बच्चे को अपने जोखिम और अपनी जिम्मेदारी पर गोद ले रहा हूं। मैंने परिवार के अन्य सदस्यों से परामर्श किया है और वे पिल्ला/बिल्ली के बच्चे को गोद लेने के लिए सहमत हो गए हैं।",font=("Arial",15,"bold")).pack()
    Label(new,text='मैं यह सुनिश्चित करूंगा कि गोद लिए गए पालतू जानवर को गोद लेने के बाद पशु चिकित्सक द्वारा पूरी तरह से जांच की जाए।',font=("Arial",15,"bold")).pack()
    Label(new,text='2) मैं यह सुनिश्चित करूंगा कि मेरे द्वारा गोद लिए गए पशु को स्वच्छ, हवादार परिसर में रखा जाए, ठीक से खिलाया जाए और नियमित व्यायाम किया जाए।',font=("Arial",15,"bold")).pack()
    Label(new,text='3) यदि मेरे द्वारा गोद लिया गया जानवर बीमार हो जाता है, तो मैं पशु चिकित्सक से परामर्श लूंगा। मैं यह भी सुनिश्चित करूंगा कि पशु चिकित्सक द्वारा बताए गए कृमिनाशक और टीकाकरण कार्यक्रम का पालन किया जाए।',font=("Arial",15,"bold")).pack()
    Label(new,text='4) मुझे इस बात की पूरी जानकारी है कि प्रारंभिक जांच के दौरान कुछ बीमारियों जैसे गैस्ट्रो-एंटराइटिस, रेबीज, डिस्टेंपर, सांस की बीमारियों आदि का निदान पशु चिकित्सक द्वारा नहीं किया जा सकता है।',font=("Arial",15,"bold")).pack()
    Label(new,text='जिस जानवर को मैंने गोद लिया है, अगर गोद लेने के कुछ दिनों के बाद उसमें किसी बीमारी के लक्षण दिखाई देते हैं, तो मैं पशु चिकित्सक से पालतू जानवर का इलाज करवाऊंगा।',font=("Arial",15,"bold")).pack()
    Label(new,text='नियम और शर्तें',font=("Arial",20,"bold")).pack()
    Label(new,text='1) मैं पालतू जानवर को हर समय कठोर परिस्थितियों में बंधा या बाहर नहीं रखूंगा या लंबे समय तक अकेला नहीं रखूंगा।',font=("Arial",15,"bold")).pack()
    Label(new,text='2) चाहे मैं नर या मादा पिल्ले/बिल्ली के बच्चे को गोद लूं, पालतू जानवरों के अच्छे स्वास्थ्य को सुनिश्चित करने के लिए मैं इसे न्यूट्रेड/स्पाय करवाऊंगा।',font=("Arial",15,"bold")).pack()
    Label(new,text='3) यदि कुत्ते के स्वामित्व को बदलने की आवश्यकता है, तो मैं केयरटेकर/पालक और टीम A.w.h को सूचित करूंगा और प्रक्रियाओं का पालन किया जाएगा।',font=("Arial",15,"bold")).pack()
    Label(new,text='4) मुझे पता है कि गोद लिए कुत्ते को छोड़ने पर पशु क्रूरता निवारण अधिनियम, 1960 के तहत जुर्माना लगेगा। टीम एनिमल्स विथ ह्यूमैनिटी ऐसा करने पर दत्तक कुत्ते के खिलाफ कानूनी कार्रवाई की मांग कर सकती है।',font=("Arial",15,"bold")).pack()
    Label(new,text='5) मैं अपनी इच्छा से गोद लिए गए पालतू पशु का नियमानुसार नगर निगम में पंजीकरण कराऊंगा।',font=("Arial",15,"bold")).pack()
    Label(new,text='6) मुझे पता है कि उचित परामर्श और परामर्श के साथ पालतू जानवर को गोद लेने की लागत शून्य रुपये (रु. 0/-) है।',font=("Arial",15,"bold")).pack()
    Label(new,text='7) मैं मानता हूं कि गोद लिए गए पालतू जानवर की जीवन भर की अवधि में, टीम एनिमल्स विद ह्यूमैनिटी और केयरटेकर/पालक को यह निरीक्षण करने का अधिकार है कि मैं जानवर को घर पर कैसे ',font=("Arial",15,"bold")).pack()
    Label(new,text='रख रहा हूं, साथ ही मैं अपने पालतू जानवर को अन्य जगहों पर डिजिटल रूप से ले जा रहा हूं। या बिना किसी पूर्व सूचना के भौतिक मुलाकातों द्वारा।',font=("Arial",15,"bold")).pack()
    Label(new,text='मैं टीम एनिमल्स विद ह्यूमैनिटी को गोद लिए गए पालतू जानवरों के बारे में नियमित अपडेट दूंगा। यदि क्रूरता निवारण अधिनियम, 1960 के तहत कोई उल्लंघन होता है, तो मुझे पता है कि टीम एनिमल्स विथ ह्यूमैनिटी जानवर को जब्त कर',font=("Arial",15,"bold")).pack()
    Label(new,text='सकती है या कानून के अनुसार कार्रवाई कर सकती है।',font=("Arial",15,"bold")).pack()
    Label(new,text='यदि पिल्ला/बिल्ली का बच्चा/पालतू पशु मालिक की लापरवाही के कारण लावारिस, अस्वस्थ, बुरी स्थिति में या मृत पाया जाता है, तो टीम एनिमल्स ',font=("Arial",15,"bold")).pack()
    Label(new,text='विथ ह्यूमैनिटी रुपये का जुर्माना ले सकती है। 10,000/- और दत्तक ग्रहण करने वाले पर कानूनी कार्रवाई के लिए आगे बढ़ें।',font=("Arial",15,"bold")).pack()
    Label(new,text="शर्तों पर हस्ताक्षर करने के बाद, यदि मैं (दत्तक) अपनाए गए पालतू जानवर को लौटाता हूं, तो मैं (दत्तक)  5000/- रुपये तक की राशि का भुगतान करने के लिए उत्तरदायी होगा लौटाने पर।",font=("Arial",15,"bold")).pack()
    Label(new,text=" मैं (दत्तक) गोद लिए गए पालतू जानवर को केवल टीम AWH को मानवता के साथ लौटाऊंगा और इसे नहीं छोड़ूंगा।",font=("Arial",15,"bold")).pack()
    Button(new,text="ठीक",command=new.destroy).pack()
    
    
    
    



Test_file()
window=Tk()
icon=PhotoImage(file=os.path.join(os.getcwd(),'icon.png'))
window.iconphoto(False,icon)
window.title('ADOPTION FORM')



label9=Label(window,text='Counselor Name',font=("Arial",15,"bold"))
entry7=Entry(window)
label1=Label(window,text='Description of the Animal',font=("Arial",15,"bold"))
label2=Label(window,text='Kind of Animal',font=("Arial",15,"bold"))
entry1=Entry(window)
label28=Label(window,text='Tag number of animal',font=("Arial",15,"bold"))
entry22=Entry(window)
label3=Label(window,text='Sex',font=("Arial",15,"bold"))
entry2=Entry(window)
label4=Label(window,text='Age',font=("Arial",15,"bold"))
entry3=Entry(window)
label5=Label(window,text='Fur color',font=("Arial",15,"bold"))
entry4=Entry(window)
label6=Label(window,text='Breed',font=("Arial",15,"bold"))
entry5=Entry(window)
label7=Label(window,text='Health acknowledgement',font=("Arial",15,"bold"))
entry6=Entry(window,width=20)
Label(window,text='Specifics of the Caretaker/Foster',font=("Arial",15,"bold")).grid(row=7,column=1)
label10=Label(window,text='Caretaker/Foster’s name',font=("Arial",15,"bold"))
entry8=Entry(window)
label11=Label(window,text='Contact no.',font=("Arial",15,"bold"))
entry9=Entry(window)
label12=Label(window,text='Whatsapp number',font=("Arial",15,"bold"))
entry10=Entry(window)
label13=Label(window,text='Email',font=("Arial",15,"bold"))
entry11=Entry(window)
label14=Label(window,text='Local Residence',font=("Arial",15,"bold"))
entry12=Entry(window)
label15=Label(window,text='Permanent residence',font=("Arial",15,"bold"))
entry13=Entry(window)
Label(window,text='Specifics of the Adopter',font=("Arial",15,"bold")).grid(sticky="W",row=13+1,column=1)
label16=Label(window,text='Adopter\'s name',font=("Arial",15,"bold"))
entry14=Entry(window)
label17=Label(window,text='Contact no',font=("Arial",15,"bold"))
entry15=Entry(window)
label18=Label(window,text='whatsapp number',font=("Arial",15,"bold"))
entry16=Entry(window)
label19=Label(window,text='Line of work',font=("Arial",15,"bold"))
entry17=Entry(window)
label20=Label(window,text='Email',font=("Arial",15,"bold"))
entry18=Entry(window)
label21=Label(window,text='Local Residence',font=("Arial",15,"bold"))
entry19=Entry(window)
label22=Label(window,text='Permanent Residence',font=("Arial",15,"bold"))
entry20=Entry(window)
Label(window,text='What are your plans for taking care of your',font=("Arial",15,"bold")).grid(sticky="e",row=20+1,column=0)
label23=Label(window,text='pet in case you go out for some days?',font=("Arial",15,"bold"))
entry21=Entry(window)
label24=Label(window,text='Have you pet a pet before',font=("Arial",15,"bold"))
menu2=StringVar()
menu2.set('please choose one')
drop2 = OptionMenu(window,menu2,*['currently own','yes','no'] )
label25=Label(window,text='Your dog will primarily be an',font=("Arial",15,"bold"))
menu3=StringVar()
menu3.set('please choose one')
drop3=OptionMenu(window,menu3,*['inside pet','gaurd pet'])
label26=Label(window,text='My dog/cat needs to be able to be alone (per day)',font=("Arial",15,"bold"))
menu4=StringVar()
menu4.set('please choose one')
drop4=OptionMenu(window,menu4,*['2 hours or less','4 hours or less','4-8 hours','8-10 hours','12 hours'])
label27=Label(window,text='When you\'re not at home, your dog/cat will spend his/her time',font=("Arial",15,"bold"))
menu5=StringVar()
menu5.set('please choose one')
drop5=OptionMenu(window,menu5,*['With the relatives','In a crate in the house','In the yard','Loose in the house','Confined to one room in the house'])
button=Button(window,text="save,print and clear",command=Save)
	
#menu6=StringVar()
#menu6.set('choose one')
#drop6=OptionMenu(window,menu6,*['yes','no'])

label1.grid(sticky="W",row=0+1,column=1,columnspan=2)
label2.grid(sticky="e",row=1+1,column=0)
entry1.grid(sticky="W",row=1+1,column=1)
label28.grid(sticky="e",row=3,column=0)
entry2.grid(sticky="W",row=3,column=1)
label3.grid(sticky="e",row=2+1+1,column=0)
entry22.grid(sticky="W",row=2+1+1,column=1)
label4.grid(sticky="e",row=2+1+1,column=2)
entry3.grid(sticky="W",row=2+1+1,column=3)
label5.grid(sticky="e",row=3+1+1,column=0)
entry4.grid(sticky="W",row=3+1+1,column=1)
label6.grid(sticky="e",row=3+1+1,column=2)
entry5.grid(sticky="W",row=3+1+1,column=3)
label7.grid(sticky="e",row=4+1+1,column=0)
#drop1.grid(sticky="W",row=4,column=1)
#label8.grid(sticky="W",row=5,column=0)
entry6.grid(sticky="W",row=4+1+1,column=1)
label9.grid(sticky="e",row=0,column=0)
entry7.grid(sticky="W",row=0,column=1)
label10.grid(sticky="e",row=8+1,column=0)
entry8.grid(sticky="W",row=8+1,column=1)
label11.grid(sticky="e",row=9+1,column=0)
entry9.grid(sticky="W",row=9+1,column=1)
label12.grid(sticky="e",row=9+1,column=2)
entry10.grid(sticky="W",row=9+1,column=3)
label13.grid(sticky="e",row=10+1,column=0)
entry11.grid(sticky="W",row=10+1,column=1)
label14.grid(sticky="e",row=11+1,column=0)
entry12.grid(sticky="W",row=11+1,column=1)
label15.grid(sticky="e",row=12+1,column=0)
entry13.grid(sticky="W",row=12+1,column=1)
label16.grid(sticky="e",row=14+1,column=0)
entry14.grid(sticky="W",row=14+1,column=1)
label17.grid(sticky="e",row=15+1,column=0)
entry15.grid(sticky="W",row=15+1,column=1)
label18.grid(sticky="e",row=15+1,column=2)
entry16.grid(sticky="W",row=15+1,column=3)
label19.grid(sticky="e",row=16+1,column=0)
entry17.grid(sticky="W",row=16+1,column=1)
label20.grid(sticky="e",row=17+1,column=0)
entry18.grid(sticky="W",row=17+1,column=1)
label21.grid(sticky="e",row=18+1,column=0)
entry19.grid(sticky="W",row=18+1,column=1)
label22.grid(sticky="e",row=19+1,column=0)
entry20.grid(sticky="W",row=19+1,column=1)
label23.grid(sticky="e",row=21+1,column=0)
entry21.grid(sticky="W",row=21+1,column=1)
label24.grid(sticky="e",row=23+1,column=0)
drop2.grid(sticky="W",row=23+1,column=1)
label25.grid(sticky="e",row=24+1,column=0)
drop3.grid(sticky="W",row=24+1,column=1)
label26.grid(sticky="e",row=25+1,column=0)
drop4.grid(sticky="W",row=25+1,column=1)
label27.grid(sticky="e",row=26+1,column=0)
drop5.grid(sticky="W",row=26+1,column=1)
button.grid(row=29,column=1)
#drop6.grid(sticky="W",row=25,column=1)


Terms_window=Toplevel()
Terms_window.iconphoto(False,icon)
Terms_window.title('Terms and Conditions and Consent form')
Label(Terms_window,text='CONSENT BY ADOPTER',font=("Arial", 25,"bold")).pack()
Label(Terms_window,text='1) I am adopting the puppy/kitten at my own risk and my responsibility. I have consulted other members of the family and they have agreed to adopt the pup/kitten.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='I will make sure that the adopted pet gets checked thoroughly by the Vet after adoption.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='2) I will ensure that the animal adopted by me is kept in clean, well-ventilated premises, properly fed and given regular exercise.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='3)In case the animal adopted by me falls ill, I will consult a vet. I will also ensure that the deworming and the vaccination schedule advised by the vet is followed.',font=("Arial",15,"bold")).pack()	
Label(Terms_window,text='4) I am fully aware that some diseases like gastro-enteritis, rabies, distemper, respiratory diseases, etc. cannot be diagnosed by a vet during the preliminary checkup.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='If the animal that I have adopted shows signs of some illness after a few days of adoption, then I will get the pet treated from a vet.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='Terms and Conditions',font=('Arial',25,"bold")).pack()
#Label(Terms_window,text='1) I shall not hold the organizers of the adoption camp responsible if the animal shows sign of illness or dies after adoption.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='1) 1 will not keep the pet tied or outside in harsh conditions all the time or all by itself for long stretches of time.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='2) Whether I adopt a male or a female puppy/kitten, I will get it neutered/ spayed to ensure good health of the pet.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='3) In case the ownership of the dog needs to be changed, I will inform the CARETAKERS/FOSTERS and Team Animals with Humanity and the procedures will be followed.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='4) I am aware that Abandoning an adopted dog will attract penalty under Prevention of Cruelty to Animals Act, 1960. Team Animals With Humanity can seek legal actions against the Adopter if done so.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='5) I will register the pet I adopted at my will with the Municipal Corporation as per the rules. ',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='6) I\'m aware that the pet adoption is at Zero Rupees (Rs. 0/-) cost with proper counseling & consultation.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='7) I agree that over the lifetime period of the Adopted Pet, Team Animals With Humanity and the Caretaker/Foster have the right to inspect the manner in which I am keeping the animal at home',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='along with other places I take my pet to, digitally or by physical meet ups WITHOUT PRIOR NOTICE',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='I will give Team Animals With Humanity regular updates of the Adopted Pet. If there are any violations under the Prevention of Cruelty Act, 1960,',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='I am aware that Team Animals With Humanity may confiscate the animal or take action as per the law.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='In case the puppy/kitten/pet is found abandoned, in a unhealthy, bad condition or dead due to negligence of the owner,',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='Team Animals with Humanity can take a fine of Rs. 10,000/- and proceed with legal action on the Adopter',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='After signing to the terms, in case, I (adopter) return the Adopted Pet, I (adopter) will be liable to pay an amount up to Rs. 5000/- if returned. ',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='I (adopter) will return the Adopted Pet ONLY to Team Animals with Humanity and won\'t abandon it.',font=("Arial",15,"bold")).pack()
Label(Terms_window,text='I enter into this contract of my own free will and understand that this is a binding contract enforceable by civil law.',font=("Arial",15,"bold")).pack()
check=Checkbutton(Terms_window,text='I consent to the above terms',font=("Arial",15,"bold"))
check.pack()
Button(Terms_window,text='translate to hindi',command=translate).pack()
Terms_window.mainloop()
